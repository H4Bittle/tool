import os
import json
import io
import uuid
from datetime import datetime

from docx.shared import Inches, Pt
from openpyxl import load_workbook
from backend.models import load_applications, load_vulnerabilities
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn

# --- Rich text for Excel (robust imports) ---
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont as XLInlineFont
from openpyxl.styles import Alignment

from PIL import Image, ImageChops, ImageOps, ImageDraw

# ------------------------- CONFIG TOGGLES -------------------------
# Flip these without changing any logic below.

BAKE_BITMAP_BORDER      = True    # safest: draws the rectangle into pixels
BITMAP_BORDER_MODE      = "inset" # "inset" (slimmer) or "expand" (classic frame)
BITMAP_BORDER_PX        = 1       # minimum 1 px (hairline). 2 ≈ ~0.5 pt
BORDER_COLOR_RGB        = (128, 128, 128)   # try (64,64,64) for softer gray
AA_FADE_STRENGTH        = 0.5     # 0.0 (no fade) .. 1.0 (very light inner line)

ADD_PICTURE_OUTLINE_XML = False   # optional: vector outline (fractions like 0.25 pt)
PICTURE_OUTLINE_PT      = 0.25    # used only if ADD_PICTURE_OUTLINE_XML = True

ADD_PARAGRAPH_BORDER    = False   # optional: border around the image paragraph
# ------------------------------------------------------------------

LOG_FILE = os.path.join(os.path.dirname(__file__), 'data', 'audit_logs.json')
WORD_TEMPLATE = os.path.join(os.path.dirname(__file__), 'templates', 'report_template.docx')
EXCEL_TEMPLATE = os.path.join(os.path.dirname(__file__), 'templates', 'excel_template.xlsx')
DOWNLOAD_DIR = os.path.join(os.path.expanduser("~"), "Downloads")
SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), '..', 'static', 'screenshots')

TMP_IMG_DIR = os.path.join(os.path.dirname(__file__), '.export_image_cache')
os.makedirs(DOWNLOAD_DIR, exist_ok=True)
os.makedirs(TMP_IMG_DIR, exist_ok=True)

def log_action(message):
    log_entry = {"timestamp": datetime.utcnow().isoformat(), "action": message}
    logs = []
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, 'r') as f:
            logs = json.load(f)
    logs.append(log_entry)
    with open(LOG_FILE, 'w') as f:
        json.dump(logs, f, indent=2)

def get_color(severity):
    return {
        "critical": "C00000", "high": "EE0000", "medium": "FFC000",
        "low": "00B050", "info": "0070C0"
    }.get(severity.lower(), "000000")

def format_human_date(date_str):
    for fmt in ("%Y-%m-%d", "%d-%m-%Y"):
        try:
            dt = datetime.strptime(date_str, fmt); break
        except Exception:
            continue
    else:
        return date_str
    day = dt.day
    suffix = 'th' if 11 <= day <= 13 else {1:'st',2:'nd',3:'rd'}.get(day % 10, 'th')
    return f"{day}{suffix} {dt.strftime('%B %Y')}"

# ----------------------------- HELPERS -----------------------------

def _resolve_screenshot_path(raw_value: str) -> str | None:
    if not raw_value:
        return None
    p = raw_value.strip().strip('"').strip("'")
    if os.path.isabs(p) and os.path.isfile(p):
        return p
    for c in (
        os.path.join(SCREENSHOT_DIR, os.path.basename(p)),
        os.path.join(SCREENSHOT_DIR, p),
    ):
        c = os.path.normpath(c)
        if os.path.isfile(c):
            return c
    return None

def _png_temp_path():
    return os.path.join(TMP_IMG_DIR, f"{uuid.uuid4().hex}.png")

def _trim_transparent_or_white(im: Image.Image, white_tol: int = 10) -> Image.Image:
    """
    Trim fully transparent edges (RGBA) and near-white matte (RGB).
    """
    try:
        if im.mode in ("RGBA", "LA"):
            alpha = im.split()[-1]
            bbox = alpha.getbbox()
            if bbox:
                im = im.crop(bbox)
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        bg = Image.new("RGB", im.size, (255, 255, 255))
        diff = ImageChops.difference(im.convert("RGB"), bg)
        enhance = ImageChops.add(diff, diff, 2.0, -white_tol)
        bbox2 = enhance.getbbox()
        if bbox2:
            im = im.crop(bbox2)
        return im
    except Exception:
        return im

def _add_bitmap_border_inset(im: Image.Image, border_px: int = 1,
                             color=(0, 0, 0), fade_strength: float = 0.5) -> Image.Image:
    """
    Draw a slim inset border INSIDE the image bounds with anti-aliased inside fade.
    - Solid outer 1 px line
    - Inner 1 px (or more) line(s) that softly blend toward white
    This reads thinner than a hard 1 px frame but stays visible on all sides.
    """
    border_px = max(1, int(border_px))
    w, h = im.size
    if w < 2*border_px + 2 or h < 2*border_px + 2:
        return ImageOps.expand(im.convert("RGB"), border=border_px, fill=color)

    if im.mode != "RGBA":
        im = im.convert("RGBA")

    inner_w, inner_h = w - 2*border_px, h - 2*border_px
    content = im.resize((inner_w, inner_h), Image.LANCZOS)

    canvas = Image.new("RGBA", (w, h), (255, 255, 255, 255))
    canvas.paste(content, (border_px, border_px))

    draw = ImageDraw.Draw(canvas)
    solid_col = (*color, 255)

    def faded(t: float):
        t = max(0.0, min(1.0, t))
        r, g, b = color
        rr = int((1 - t) * r + t * 255)
        gg = int((1 - t) * g + t * 255)
        bb = int((1 - t) * b + t * 255)
        return (rr, gg, bb, 255)

    draw.rectangle([0, 0, w-1, h-1], outline=solid_col, width=1)

    for i in range(1, border_px):
        t = fade_strength * (i / border_px)
        fade_col = faded(t)
        draw.rectangle([i, i, w-1-i, h-1-i], outline=fade_col, width=1)

    return canvas.convert("RGB")

def _add_bitmap_border_expand(im: Image.Image, border_px: int = 1, color=(0, 0, 0)) -> Image.Image:
    border_px = max(1, int(border_px))
    return ImageOps.expand(im.convert("RGB"), border=border_px, fill=color)

def _inline_image_force_png_path(doc_tpl, abs_path: str, width_in: float = 6.48):
    """
    Load with Pillow, trim edges, add bitmap border (inset/expand), save to PNG on disk,
    and return InlineImage(path,...). File path is most stable for docxtpl/python-docx.
    """
    out_path = _png_temp_path()
    with Image.open(abs_path) as im:
        im.load()
        im = _trim_transparent_or_white(im)
        if BAKE_BITMAP_BORDER:
            if BITMAP_BORDER_MODE.lower() == "inset":
                im = _add_bitmap_border_inset(
                    im, border_px=BITMAP_BORDER_PX, color=BORDER_COLOR_RGB, fade_strength=AA_FADE_STRENGTH
                )
            else:
                im = _add_bitmap_border_expand(im, border_px=BITMAP_BORDER_PX, color=BORDER_COLOR_RGB)
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        im.save(out_path, format="PNG")
    return InlineImage(doc_tpl, out_path, width=Inches(width_in))

def _set_paragraph_border(paragraph, size_pt=0.5, color_hex="000000"):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag == qn('w:pBdr'):
            pPr.remove(child)
    pBdr = OxmlElement('w:pBdr')
    sz = str(int(size_pt * 8))  # Word uses 1/8 pt units
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        pBdr.append(el)
    pPr.append(pBdr)

def _style_image_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if ADD_PARAGRAPH_BORDER:
        _set_paragraph_border(paragraph, size_pt=0.5, color_hex="000000")

def _postprocess_outline_and_layout(docx_path: str):
    """
    After rendering:
      - center-align image paragraphs and zero their spacing
      - optionally add picture outline via DrawingML
    """
    d = Document(docx_path)

    for p in d.paragraphs:
        if p._element.xpath('.//w:drawing'):
            _style_image_paragraph(p)

    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p._element.xpath('.//w:drawing'):
                        _style_image_paragraph(p)

    for sec in d.sections:
        for container in (sec.header, sec.footer):
            if container:
                for p in container.paragraphs:
                    if p._element.xpath('.//w:drawing'):
                        _style_image_paragraph(p)

    if ADD_PICTURE_OUTLINE_XML:
        _add_picture_outline_xml(d, PICTURE_OUTLINE_PT)

    d.save(docx_path)

def _add_picture_outline_xml(doc: Document, line_pt: float = 0.5):
    """
    Add a solid outline to each picture via DrawingML.
    Toggle with ADD_PICTURE_OUTLINE_XML.
    """
    roots = [doc.element.body]
    for sec in doc.sections:
        if sec.header: roots.append(sec.header._element)
        if sec.footer: roots.append(sec.footer._element)

    for root in roots:
        if root is None:
            continue
        pics = root.xpath('.//w:drawing//pic:pic')
        for pic in pics:
            spPr_list = pic.xpath('./pic:spPr')
            if spPr_list:
                spPr = spPr_list[0]
            else:
                spPr = parse_xml(
                    '<pic:spPr xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
                    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>'
                )
                pic.append(spPr)
            for ln in spPr.xpath('./a:ln'):
                spPr.remove(ln)
            w_emu = int(12700 * max(0.1, float(line_pt)))
            ln_xml = (
                f'<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="{w_emu}">'
                f'  <a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
                f'</a:ln>'
            )
            spPr.append(parse_xml(ln_xml))

# ----------------------------- MAIN API -----------------------------

def generate_word_report(app_id):
    template_path = WORD_TEMPLATE
    app_file = os.path.join(os.path.dirname(__file__), 'data', 'applications', f'{app_id}.json')
    vuln_file = os.path.join(os.path.dirname(__file__), 'data', 'vulnerabilities', f'{app_id}.json')
    if not all(os.path.exists(p) for p in [template_path, app_file, vuln_file]):
        return None

    doc = DocxTemplate(template_path)
    with open(app_file) as f: app_data = json.load(f)
    with open(vuln_file) as f: vulnerabilities = json.load(f)
    vulnerabilities.sort(key=lambda x: float(x.get("cvss", 0)), reverse=True)

    vuln_summary, vuln_details = [], []

    for vuln in vulnerabilities:
        vuln_summary.append({
            "vulnerability_id": vuln.get("id", ""),
            "title": vuln.get("title", ""),
            "cvss_score": str(vuln.get("cvss", "")),
            "severity": vuln.get("severity", "")
        })
        step_entries = []
        for idx, step in enumerate(vuln.get("steps", []), 1):
            image_filename = (step.get("screenshot") or "").strip()
            resolved_path = _resolve_screenshot_path(image_filename)
            exists = os.path.exists(resolved_path) if resolved_path else False
            is_valid = False
            if exists and image_filename:
                try:
                    with Image.open(resolved_path) as im:
                        im.verify()
                    is_valid = True
                except Exception as ex:
                    print(f"Invalid image file: {resolved_path} | Error: {ex}")
            else:
                print(f"Image file missing or no filename: {os.path.join(SCREENSHOT_DIR, image_filename)}")

            print(f"Screenshot for step {idx}: {resolved_path}, exists: {exists}, valid: {is_valid}, screenshot val: {image_filename}")
            screenshot_obj = ""
            if is_valid and resolved_path:
                try:
                    screenshot_obj = _inline_image_force_png_path(doc, resolved_path, width_in=6.48)
                except Exception as ex:
                    print(f"Failed to embed screenshot for {resolved_path}: {ex}")

            step_entries.append({
                "index": idx,
                "description": step.get("description", ""),
                "screenshot": screenshot_obj
            })

        vuln_details.append({
            "title": vuln.get("title", ""),
            "vulnerability_id": vuln.get("id", ""),
            "summary": vuln.get("summary", ""),
            "description": vuln.get("description", ""),
            "business_impact": vuln.get("impact", ""),
            "severity": vuln.get("severity", ""),
            "cvss_score": str(vuln.get("cvss", "")),
            "cvss_vector": str(vuln.get("cvss_vector", "")),
            "affected_url": vuln.get("url", ""),
            "recommendation": vuln.get("recommendation", ""),
            "cwe_id": vuln.get("cwe", ""),
            "reference": vuln.get("reference", ""),
            "step_entries": step_entries
        })

    app_details = [{
        "index": i+1, "app_name": a.get("name", ""), "app_version": a.get("version", ""), "app_url": a.get("url", "")
    } for i, a in enumerate(app_data.get("app_details", []))]

    pentesters = [{
        "pentester_name": p.get("name", ""), "pentester_role": p.get("role", ""), "pentester_email": p.get("email", "")
    } for p in app_data.get("pentesters", [])]

    context = {
        "app_name": app_data.get("name", ""),
        "start_date": format_human_date(app_data.get("start_date", "")),
        "end_date": format_human_date(app_data.get("end_date", "")),
        "vulnerabilities": vuln_summary,
        "vuln_details": vuln_details,
        "app_details": app_details,
        "pentesters": pentesters,
        "test_credentials": app_data.get("test_credentials", [])
    }

    safe_name = app_data.get("name", "Report").replace("/", "_").replace("\\", "_").replace(" ", "_")
    filename = f"GW_{safe_name}_Penetration_Test_Report.docx"
    output_path = os.path.join(DOWNLOAD_DIR, filename)

    print("START DATE:", context["start_date"])
    print("END DATE:", context["end_date"])
    doc.render(context)
    doc.save(output_path)

    _postprocess_outline_and_layout(output_path)

    d = Document(output_path)
    for table in d.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if "risk rating" in headers: idx = headers.index("risk rating")
        elif "severity" in headers: idx = headers.index("severity")
        else: continue
        for row in table.rows[1:]:
            sev = row.cells[idx].text.strip().lower()
            if sev in ["critical", "high", "medium", "low", "info"]:
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_pr.append(parse_xml(
                    f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    f'w:val="clear" w:color="auto" w:fill="{get_color(sev)}"/>'
                ))

    for table in d.tables:
        for row in table.rows:
            if len(row.cells) >= 2 and row.cells[0].text.strip().lower() == "severity":
                sev = row.cells[1].text.strip().lower()
                if sev in ["critical", "high", "medium", "low", "info"]:
                    tc_pr = row.cells[1]._tc.get_or_add_tcPr()
                    tc_pr.append(parse_xml(
                        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        f'w:val="clear" w:color="auto" w:fill="{get_color(sev)}"/>'
                    ))

    d.save(output_path)
    return output_path

def generate_excel_report(app_id):
    # Load the Excel template
    wb = load_workbook(EXCEL_TEMPLATE)
    ws = wb.active

    # Get app and its vulnerabilities
    app = next((a for a in load_applications() if a['id'] == app_id), None)
    if not app:
        return None

    vulns = load_vulnerabilities(app_id)

    # Sort by highest CVSS first (descending)
    vulns_sorted = sorted(
        vulns,
        key=lambda v: float(v.get("cvss", 0) or 0),
        reverse=True
    )

    # helper for rich text blocks (bold or normal)
    def tb(text, bold=False):
        return TextBlock(text=text, font=XLInlineFont(b=bool(bold)))

    # Write rows to the sheet
    for v in vulns_sorted:
        redacted_summary = v.get("summary", "") or ""
        long_description = v.get("description", "") or ""
        steps_block = "\n".join(
            [f"Step {i+1}: {s.get('description','')}" for i, s in enumerate(v.get("steps", []))]
        )

        # Append a row. Put None in the Description and Finding Tag columns for now.
        ws.append([
            app.get("name", ""),            # A: App Name
            v.get("url", ""),               # B: Affected URL
            v.get("id", ""),                # C: Vulnerability ID
            v.get("cvss", 0),               # D: CVSS Score (numeric)
            "",                             # E: placeholder (match template)
            v.get("title", ""),             # F: Title
            None,                           # G: Description (rich text inserted below)
            v.get("impact", ""),            # H: Business Impact
            "",                             # I: placeholder
            v.get("recommendation", ""),    # J: Recommendation
            v.get("reference", ""),         # K: Reference
            None                            # L: Finding Tag (formula set below)
        ])

        # Build rich text for the Description cell (column 7 → "G")
        rt = CellRichText()
        rt.append(tb("Redacted Summary:\n", bold=True))
        rt.append(tb(f"{redacted_summary}\n\n", bold=False))

        rt.append(tb("Description:\n", bold=True))
        rt.append(tb(f"{long_description}\n", bold=False))

        if steps_block.strip():
            rt.append(tb("\nSteps to Reproduce:\n", bold=True))
            rt.append(tb(steps_block, bold=False))

        last_row = ws.max_row

        # Set Description rich text
        desc_cell = ws.cell(row=last_row, column=7)  # G
        desc_cell.value = rt
        desc_cell.alignment = Alignment(wrap_text=True)

        # Set Finding Tag formula in column L (12)
        ws.cell(
     row=last_row,
   column=12
).value = f'=A{last_row}&", "&C{last_row}&", "&F{last_row}'

    # Build filename: GW_<App_Name>_Excel_Findings_<Count>.xlsx
    safe_name = (app.get("name", "Report")
                 .replace("/", "_").replace("\\", "_").replace(" ", "_"))
    findings_count = len(vulns_sorted)
    filename = f"GW_{safe_name}_Excel_Findings_{findings_count}.xlsx"
    output_path = os.path.join(DOWNLOAD_DIR, filename)

    wb.save(output_path)
    return output_path
